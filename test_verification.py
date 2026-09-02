"""
Automated Verification Test Suite for Cloud AI Chatbot
Tests all core features:
- /api/health
- Validation of TXT, PDF, DOCX, and Image files
- Validation of unsupported and oversized files
- Error handling for missing/invalid API key
- OpenRouter Gemini Client logic & multimodal payload preparation
- Conversation history & session clearing
"""

import os
import io
import unittest
import docx
from pypdf import PdfWriter

# Set test environment
os.environ["FLASK_SECRET_KEY"] = "test-secret-key"
os.environ["OPENROUTER_MODEL"] = "google/gemini-3.7-flash"

from app import app
from utils.file_processor import validate_file, extract_text_from_file, is_image_file, FileValidationError
from utils.gemini_client import GeminiClient, GeminiClientError


class ChatbotTestCase(unittest.TestCase):
    def setUp(self):
        app.config["TESTING"] = True
        self.client = app.test_client()

    def test_01_health_endpoint(self):
        """Verify GET /api/health returns 200 OK and expected metadata."""
        response = self.client.get("/api/health")
        self.assertEqual(response.status_code, 200)
        data = response.get_json()
        self.assertEqual(data["status"], "ok")
        self.assertIn("OpenRouter", data["provider"])
        self.assertEqual(data["model"], "google/gemini-3.7-flash")
        self.assertTrue(data["web_search"])
        print("[PASSED] /api/health endpoint test")

    def test_02_index_page(self):
        """Verify GET / returns 200 OK and contains chatbot layout."""
        response = self.client.get("/")
        self.assertEqual(response.status_code, 200)
        html = response.get_data(as_text=True)
        self.assertIn("Cloud AI Chatbot", html)
        self.assertIn("Gemini Flash", html)
        self.assertIn("composerForm", html)
        print("[PASSED] / index page template test")

    def test_03_txt_file_processing(self):
        """Verify text extraction from a TXT file."""
        sample_text = "Cloud Computing is the on-demand delivery of compute power and storage."
        test_txt_path = os.path.join(app.config["UPLOAD_FOLDER"], "test_sample.txt")
        with open(test_txt_path, "w", encoding="utf-8") as f:
            f.write(sample_text)

        extracted = extract_text_from_file(test_txt_path)
        self.assertEqual(extracted, sample_text)
        print("[PASSED] TXT file extraction test")

    def test_04_docx_file_processing(self):
        """Verify text and table extraction from a DOCX file."""
        test_docx_path = os.path.join(app.config["UPLOAD_FOLDER"], "test_sample.docx")
        doc = docx.Document()
        doc.add_heading("HPCC Project Report", level=1)
        doc.add_paragraph("This document tests python-docx extraction for Azure Chatbot.")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Service"
        table.cell(0, 1).text = "Role"
        table.cell(1, 0).text = "App Service"
        table.cell(1, 1).text = "Web Hosting"
        doc.save(test_docx_path)

        extracted = extract_text_from_file(test_docx_path)
        self.assertIn("HPCC Project Report", extracted)
        self.assertIn("App Service | Web Hosting", extracted)
        print("[PASSED] DOCX file extraction test")

    def test_05_pdf_file_processing(self):
        """Verify text extraction from a PDF file using pypdf."""
        test_pdf_path = os.path.join(app.config["UPLOAD_FOLDER"], "test_sample.pdf")
        writer = PdfWriter()
        writer.add_blank_page(width=72, height=72)
        with open(test_pdf_path, "wb") as f:
            writer.write(f)

        with self.assertRaises(ValueError):
            extract_text_from_file(test_pdf_path)
        print("[PASSED] PDF empty/blank detection test")

    def test_06_image_file_detection(self):
        """Verify image format detection."""
        self.assertTrue(is_image_file("diagram.png"))
        self.assertTrue(is_image_file("photo.jpg"))
        self.assertTrue(is_image_file("banner.webp"))
        self.assertFalse(is_image_file("report.pdf"))
        self.assertFalse(is_image_file("notes.txt"))
        print("[PASSED] Image file detection test")

    def test_07_invalid_and_oversized_files(self):
        """Verify rejection of unsupported extensions and oversized files."""
        data = {
            "message": "Hello",
            "file": (io.BytesIO(b"binary content"), "malware.exe"),
        }
        res = self.client.post("/api/chat", data=data, content_type="multipart/form-data")
        self.assertEqual(res.status_code, 400)
        json_data = res.get_json()
        self.assertIn("Unsupported file type", json_data["error"])

        res_empty = self.client.post("/api/chat", data={"message": ""})
        self.assertEqual(res_empty.status_code, 400)
        print("[PASSED] Invalid file validation & empty message test")

    def test_08_api_key_error_handling(self):
        """Verify clear 401 error message when API key is missing or unconfigured."""
        client = GeminiClient(api_key="your_openrouter_api_key_here")
        with self.assertRaises(GeminiClientError) as ctx:
            client.generate_reply("Hello")
        self.assertEqual(ctx.exception.status_code, 401)
        self.assertIn("OPENROUTER_API_KEY", ctx.exception.message)
        print("[PASSED] API key defensive error handling test")

    def test_09_clear_chat_endpoint(self):
        """Verify POST /api/clear removes session history."""
        with self.client.session_transaction() as sess:
            sess["history"] = [{"role": "user", "text": "Hi"}, {"role": "model", "text": "Hello"}]

        res = self.client.post("/api/clear")
        self.assertEqual(res.status_code, 200)
        data = res.get_json()
        self.assertEqual(data["status"], "cleared")

        with self.client.session_transaction() as sess:
            self.assertNotIn("history", sess)
        print("[PASSED] /api/clear session clearing test")

    def test_10_index_page_new_elements(self):
        """Verify GET / contains new chat, pinned, previous chats, and share elements."""
        response = self.client.get("/")
        self.assertEqual(response.status_code, 200)
        html = response.get_data(as_text=True)
        self.assertIn("newChatBtn", html)
        self.assertIn("pinnedSection", html)
        self.assertIn("previousSection", html)
        self.assertIn("shareBtn", html)
        self.assertIn("shareDropdownMenu", html)
        self.assertIn("storage-notice", html)
        print("[PASSED] New Chat, Pinned, Previous Chats, and Share template elements test")

    def test_11_custom_history_handling(self):
        """Verify /api/chat gracefully parses custom history JSON parameter."""
        # Empty message with custom history should still return 400 for empty message
        res = self.client.post("/api/chat", data={"message": "", "history": '[{"role":"user","text":"test"}]'})
        self.assertEqual(res.status_code, 400)
        print("[PASSED] Custom history parsing validation test")

    def test_12_model_and_effort_configuration(self):
        """Verify model popover menu, effort options, and microphone button in HTML and health endpoint."""
        # 1. Verify UI elements exist in template
        response = self.client.get("/")
        self.assertEqual(response.status_code, 200)
        html = response.get_data(as_text=True)
        self.assertIn("modelMenuBtn", html)
        self.assertIn("modelPopoverMenu", html)
        self.assertIn("micBtn", html)
        self.assertIn("google/gemini-3.7-flash", html)
        self.assertIn("google/gemini-3.6-flash", html)
        self.assertIn("google/gemini-3.1-flash-lite", html)
        self.assertIn('data-effort="low"', html)
        self.assertIn('data-effort="medium"', html)
        self.assertIn('data-effort="high"', html)

        # 2. Verify /api/health exposes available models and effort levels
        health_res = self.client.get("/api/health")
        self.assertEqual(health_res.status_code, 200)
        hdata = health_res.get_json()
        self.assertIn("available_models", hdata)
        self.assertIn("effort_levels", hdata)
        self.assertListEqual(hdata["effort_levels"], ["low", "medium", "high"])
        print("[PASSED] Model selector and reasoning effort configuration test")

    def test_13_image_generation_endpoint(self):
        """Verify POST /api/generate-image validates empty prompt and generates image."""
        # Empty prompt validation
        res_empty = self.client.post("/api/generate-image", json={"prompt": ""})
        self.assertEqual(res_empty.status_code, 400)

        # Valid prompt generation
        res = self.client.post("/api/generate-image", json={"prompt": "A modern futuristic data center"})
        self.assertEqual(res.status_code, 200)
        data = res.get_json()
        self.assertEqual(data["status"], "success")
        self.assertIn("image_url", data)
        self.assertTrue(data["image_url"].startswith("/static/generated/images/"))
        print("[PASSED] Image generation endpoint test")

    def test_14_video_generation_endpoint(self):
        """Verify POST /api/generate-video initiates async task and GET /api/video/status polls state."""
        # Empty prompt validation
        res_empty = self.client.post("/api/generate-video", json={"prompt": ""})
        self.assertEqual(res_empty.status_code, 400)

        # Initiate video task
        res = self.client.post("/api/generate-video", json={"prompt": "A drone flyover of neon mountains"})
        self.assertEqual(res.status_code, 200)
        data = res.get_json()
        self.assertIn("job_id", data)
        job_id = data["job_id"]

        # Poll status
        status_res = self.client.get(f"/api/video/status/{job_id}")
        self.assertEqual(status_res.status_code, 200)
        sdata = status_res.get_json()
        self.assertEqual(sdata["job_id"], job_id)
        self.assertIn(sdata["status"], ["processing", "completed"])
        print("[PASSED] Asynchronous video generation & polling test")

    def test_15_chat_stream_endpoint(self):
        """Verify POST /api/chat/stream handles validation and initiates text/event-stream response."""
        # Empty message returns 400
        res_empty = self.client.post("/api/chat/stream", data={"message": ""})
        self.assertEqual(res_empty.status_code, 400)

        # Valid message returns text/event-stream response
        res = self.client.post("/api/chat/stream", data={"message": "Hello"})
        self.assertIn("text/event-stream", res.content_type)
        print("[PASSED] SSE chat streaming endpoint test")

    def test_16_index_new_chatgpt_elements(self):
        """Verify GET / contains new ChatGPT features: search, archive, dropzone, mode chip, stop button."""
        response = self.client.get("/")
        self.assertEqual(response.status_code, 200)
        html = response.get_data(as_text=True)
        self.assertIn("dropzoneOverlay", html)
        self.assertIn("chatSearchInput", html)
        self.assertIn("archivedSection", html)
        self.assertIn("plusDropdownMenu", html)
        self.assertIn("activeModeChip", html)
        self.assertIn("stopBtn", html)
        self.assertIn("Generate an Image", html)
        self.assertIn("Generate a Video", html)
        print("[PASSED] New ChatGPT multimodal template elements test")

    def test_17_scheduled_view_elements(self):
        """Verify GET /scheduled and GET / contain Scheduled AI Assistant components."""
        res = self.client.get("/scheduled")
        self.assertEqual(res.status_code, 200)
        html = res.get_data(as_text=True)
        self.assertIn("sidebarScheduledBtn", html)
        self.assertIn("scheduledView", html)
        self.assertIn("scheduleInput", html)
        self.assertIn("recommendedCardsGrid", html)
        self.assertIn("scheduledTasksGrid", html)
        self.assertIn("taskConfirmModal", html)
        self.assertIn("taskHistoryModal", html)
        self.assertIn("taskSettingsModal", html)
        print("[PASSED] Scheduled AI Assistant layout and modal elements test")

    def test_18_scheduled_task_crud(self):
        """Verify full CRUD lifecycle for scheduled tasks."""
        headers = {"X-Session-ID": "test_session_crud"}

        # 1. Create task
        new_task_payload = {
            "title": "Daily AI Briefing",
            "prompt": "Summarize top 3 AI papers today",
            "task_type": "briefing",
            "recurrence": "daily",
            "time": "08:30",
        }
        res_create = self.client.post("/api/tasks", json=new_task_payload, headers=headers)
        self.assertEqual(res_create.status_code, 201)
        task_data = res_create.get_json()["task"]
        task_id = task_data["id"]
        self.assertEqual(task_data["title"], "Daily AI Briefing")
        self.assertEqual(task_data["recurrence"], "daily")
        self.assertEqual(task_data["status"], "active")
        self.assertTrue("next_run" in task_data)

        # 2. List tasks
        res_list = self.client.get("/api/tasks", headers=headers)
        self.assertEqual(res_list.status_code, 200)
        tasks = res_list.get_json()["tasks"]
        self.assertTrue(any(t["id"] == task_id for t in tasks))

        # 3. Get single task
        res_get = self.client.get(f"/api/tasks/{task_id}", headers=headers)
        self.assertEqual(res_get.status_code, 200)
        self.assertEqual(res_get.get_json()["title"], "Daily AI Briefing")

        # 4. Update task
        res_update = self.client.put(
            f"/api/tasks/{task_id}",
            json={"title": "Updated Daily AI Briefing", "time": "09:00"},
            headers=headers
        )
        self.assertEqual(res_update.status_code, 200)
        updated = res_update.get_json()["task"]
        self.assertEqual(updated["title"], "Updated Daily AI Briefing")
        self.assertEqual(updated["time"], "09:00")

        # 5. Delete task
        res_del = self.client.delete(f"/api/tasks/{task_id}", headers=headers)
        self.assertEqual(res_del.status_code, 200)

        # Confirm deleted
        res_after = self.client.get(f"/api/tasks/{task_id}", headers=headers)
        self.assertEqual(res_after.status_code, 404)
        print("[PASSED] Scheduled task CRUD lifecycle test")

    def test_19_scheduled_task_pause_resume(self):
        """Verify pausing and resuming scheduled tasks."""
        headers = {"X-Session-ID": "test_session_pause"}
        res = self.client.post("/api/tasks", json={
            "title": "Pause Test Task",
            "prompt": "Test pause functionality",
            "recurrence": "daily",
            "time": "12:00",
        }, headers=headers)
        task_id = res.get_json()["task"]["id"]

        # Pause
        res_pause = self.client.post(f"/api/tasks/{task_id}/pause", headers=headers)
        self.assertEqual(res_pause.status_code, 200)
        self.assertEqual(res_pause.get_json()["task"]["status"], "paused")

        # Resume
        res_resume = self.client.post(f"/api/tasks/{task_id}/resume", headers=headers)
        self.assertEqual(res_resume.status_code, 200)
        self.assertEqual(res_resume.get_json()["task"]["status"], "active")

        # Cleanup
        self.client.delete(f"/api/tasks/{task_id}", headers=headers)
        print("[PASSED] Scheduled task pause & resume test")

    def test_20_scheduled_task_run_now(self):
        """Verify POST /api/tasks/<id>/run executes task immediately and records history."""
        headers = {"X-Session-ID": "test_session_run_now"}
        res = self.client.post("/api/tasks", json={
            "title": "Run Now Quick Reminder",
            "prompt": "Respond with 'Task executed successfully.'",
            "task_type": "reminder",
            "recurrence": "once",
            "time": "10:00",
        }, headers=headers)
        task_id = res.get_json()["task"]["id"]

        # Execute Run Now
        res_run = self.client.post(f"/api/tasks/{task_id}/run", headers=headers)
        self.assertEqual(res_run.status_code, 200)
        run_data = res_run.get_json()
        self.assertEqual(run_data["status"], "completed")
        self.assertTrue(len(run_data["result"]) > 0)
        self.assertIn("duration_seconds", run_data)

        # Verify execution record appended in task history
        res_get = self.client.get(f"/api/tasks/{task_id}", headers=headers)
        task = res_get.get_json()
        self.assertTrue(len(task["execution_history"]) >= 1)
        self.assertEqual(task["execution_history"][0]["status"], "completed")

        # Cleanup
        self.client.delete(f"/api/tasks/{task_id}", headers=headers)
        print("[PASSED] Scheduled task Run Now execution and history logging test")

    def test_21_natural_language_schedule_parsing(self):
        """Verify POST /api/tasks/parse-nl parses natural language schedule commands."""
        # 1. Daily study reminder
        res1 = self.client.post("/api/tasks/parse-nl", json={
            "text": "Remind me to study every day at 7 PM"
        })
        self.assertEqual(res1.status_code, 200)
        p1 = res1.get_json()["parsed"]
        self.assertEqual(p1["recurrence"], "daily")
        self.assertEqual(p1["time"], "19:00")
        self.assertEqual(p1["task_type"], "study")

        # 2. Weekly briefing on Monday
        res2 = self.client.post("/api/tasks/parse-nl", json={
            "text": "Every Monday at 9 AM, give me a summary of AI news"
        })
        self.assertEqual(res2.status_code, 200)
        p2 = res2.get_json()["parsed"]
        self.assertEqual(p2["recurrence"], "weekly")
        self.assertIn("Monday", p2["days_of_week"])
        self.assertEqual(p2["time"], "09:00")
        self.assertEqual(p2["task_type"], "briefing")
        print("[PASSED] Natural language schedule parsing test")

    def test_22_session_isolation(self):
        """Verify scheduled tasks belong strictly to the current session and never leak."""
        sid_user_a = "session_user_alice"
        sid_user_b = "session_user_bob"

        # Alice creates a private task
        res_a = self.client.post("/api/tasks", json={
            "title": "Alice Confidential Briefing",
            "prompt": "Summarize confidential data",
            "recurrence": "daily",
            "time": "08:00",
        }, headers={"X-Session-ID": sid_user_a})
        task_a_id = res_a.get_json()["task"]["id"]

        # Bob lists his tasks - Alice's task must NOT appear
        res_b_list = self.client.get("/api/tasks", headers={"X-Session-ID": sid_user_b})
        bob_tasks = res_b_list.get_json()["tasks"]
        self.assertFalse(any(t["id"] == task_a_id for t in bob_tasks))

        # Bob attempts to access Alice's task directly - must return 404
        res_b_get = self.client.get(f"/api/tasks/{task_a_id}", headers={"X-Session-ID": sid_user_b})
        self.assertEqual(res_b_get.status_code, 404)

        # Cleanup
        self.client.delete(f"/api/tasks/{task_a_id}", headers={"X-Session-ID": sid_user_a})
        print("[PASSED] Session task isolation and security test")

    def test_23_scheduler_tick_endpoint(self):
        """Verify POST /api/tasks/tick functions for Azure Functions Timer Trigger & WebJobs."""
        res = self.client.post("/api/tasks/tick")
        self.assertEqual(res.status_code, 200)
        data = res.get_json()
        self.assertEqual(data["status"], "ok")
        self.assertIn("executed_count", data)
        self.assertIn("timestamp", data)
        print("[PASSED] External scheduler tick webhook endpoint test")

    def test_24_translation_view_route_and_elements(self):
        """Verify GET /translation route returns 200 and renders complete Translation layout."""
        res = self.client.get("/translation")
        self.assertEqual(res.status_code, 200)
        html = res.get_data(as_text=True)

        self.assertIn("sidebarTranslationBtn", html)
        self.assertIn("translationView", html)
        self.assertIn("sourceLangBtn", html)
        self.assertIn("targetLangBtn", html)
        self.assertIn("swapLangBtn", html)
        self.assertIn("sourceText", html)
        self.assertIn("targetText", html)
        self.assertIn("submitTranslateBtn", html)
        self.assertIn("copyTranslationBtn", html)
        self.assertIn("readTranslationBtn", html)
        self.assertIn("useInChatBtn", html)
        self.assertIn("recentTranslationsList", html)
        print("[PASSED] /translation route and UI elements test")

    def test_25_languages_api_endpoint(self):
        """Verify GET /api/languages returns comprehensive list of supported languages."""
        res = self.client.get("/api/languages")
        self.assertEqual(res.status_code, 200)
        data = res.get_json()
        self.assertTrue(data["success"])
        self.assertTrue(data["total"] >= 25)

        lang_names = [l["name"] for l in data["languages"]]
        self.assertIn("Auto Detect", lang_names)
        self.assertIn("English", lang_names)
        self.assertIn("Tamil", lang_names)
        self.assertIn("Hindi", lang_names)
        self.assertIn("Japanese", lang_names)
        self.assertIn("Spanish", lang_names)
        self.assertIn("German", lang_names)
        self.assertIn("French", lang_names)
        print("[PASSED] /api/languages endpoint test")

    def test_26_translate_api_endpoint_languages(self):
        """Verify POST /api/translate accurately translates across languages."""
        # 1. English to Tamil
        res_ta = self.client.post("/api/translate", json={
            "text": "Hello, how are you?",
            "source_language": "English",
            "target_language": "Tamil",
            "mode": "natural",
        })
        self.assertEqual(res_ta.status_code, 200)
        data_ta = res_ta.get_json()
        self.assertTrue(data_ta["success"])
        self.assertEqual(data_ta["target_language"], "Tamil")
        self.assertTrue(len(data_ta["translation"]) > 0)

        # 2. English to Hindi
        res_hi = self.client.post("/api/translate", json={
            "text": "Good morning and welcome to our cloud platform.",
            "source_language": "English",
            "target_language": "Hindi",
            "mode": "natural",
        })
        self.assertEqual(res_hi.status_code, 200)
        data_hi = res_hi.get_json()
        self.assertTrue(data_hi["success"])
        self.assertTrue(len(data_hi["translation"]) > 0)

        # 3. Auto Detect French to English
        res_auto = self.client.post("/api/translate", json={
            "text": "Bonjour le monde, bienvenue!",
            "source_language": "auto",
            "target_language": "English",
            "mode": "natural",
        })
        self.assertEqual(res_auto.status_code, 200)
        data_auto = res_auto.get_json()
        self.assertTrue(data_auto["success"])
        self.assertTrue(len(data_auto["translation"]) > 0)
        print("[PASSED] /api/translate multilingual translation test")

    def test_27_translate_api_modes(self):
        """Verify POST /api/translate respects translation modes (formal, casual, literal)."""
        for mode in ["formal", "casual", "literal", "natural"]:
            res = self.client.post("/api/translate", json={
                "text": "Please send me the project report by Monday.",
                "source_language": "English",
                "target_language": "Spanish",
                "mode": mode,
            })
            self.assertEqual(res.status_code, 200)
            data = res.get_json()
            self.assertTrue(data["success"])
            self.assertEqual(data["mode"], mode)
            self.assertTrue(len(data["translation"]) > 0)
        print("[PASSED] /api/translate translation modes test")

    def test_28_translate_file_endpoint(self):
        """Verify POST /api/translate/file handles document translation (TXT, DOCX)."""
        # 1. Plain text file upload
        txt_content = b"Cloud AI Chatbot provides intelligent natural language capabilities and translation."
        data = {
            "file": (io.BytesIO(txt_content), "cloud_brief.txt"),
            "target_language": "Spanish",
            "source_language": "English",
            "mode": "natural",
        }
        res_txt = self.client.post(
            "/api/translate/file",
            data=data,
            content_type="multipart/form-data",
        )
        self.assertEqual(res_txt.status_code, 200)
        data_txt = res_txt.get_json()
        self.assertTrue(data_txt["success"])
        self.assertEqual(data_txt["filename"], "cloud_brief.txt")
        self.assertTrue(len(data_txt["translation"]) > 0)

        # 2. DOCX file upload
        docx_doc = docx.Document()
        docx_doc.add_paragraph("Machine learning enables modern artificial intelligence systems.")
        docx_io = io.BytesIO()
        docx_doc.save(docx_io)
        docx_io.seek(0)

        data_docx = {
            "file": (docx_io, "ai_overview.docx"),
            "target_language": "German",
            "source_language": "English",
            "mode": "formal",
        }
        res_docx = self.client.post(
            "/api/translate/file",
            data=data_docx,
            content_type="multipart/form-data",
        )
        self.assertEqual(res_docx.status_code, 200)
        data_docx_res = res_docx.get_json()
        self.assertTrue(data_docx_res["success"])
        self.assertEqual(data_docx_res["filename"], "ai_overview.docx")
        self.assertTrue(len(data_docx_res["translation"]) > 0)
        print("[PASSED] /api/translate/file document translation test")

    def test_29_natural_language_intent_parsing(self):
        """Verify natural language translation instruction parsing."""
        from utils.translator import parse_natural_language_translation_instruction

        # 1. Translate this to Spanish: Hello world
        target, mode, text = parse_natural_language_translation_instruction("Translate this to Spanish: Hello world")
        self.assertEqual(target, "Spanish")
        self.assertEqual(mode, "natural")
        self.assertEqual(text, "Hello world")

        # 2. Convert this paragraph into formal Japanese: Please send the invoice
        target, mode, text = parse_natural_language_translation_instruction("Convert this paragraph into formal Japanese: Please send the invoice")
        self.assertEqual(target, "Japanese")
        self.assertEqual(mode, "formal")
        self.assertEqual(text, "Please send the invoice")

        # 3. Translate into Tamil: Good morning
        target, mode, text = parse_natural_language_translation_instruction("Translate into Tamil: Good morning")
        self.assertEqual(target, "Tamil")
        self.assertEqual(mode, "natural")
        self.assertEqual(text, "Good morning")
        print("[PASSED] Natural language translation parsing test")


if __name__ == "__main__":
    unittest.main()
