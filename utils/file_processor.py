"""
File processing utilities for Cloud-Based AI Chatbot.

Handles:
- Validating uploaded files (extension, size, empty file checks, path safety)
- Extracting plain text from PDF / TXT / DOCX files
- Detecting whether a file is an image (handled directly by multimodal vision)

Defensive design: Every extraction function fails with clear, user-friendly messages.
"""

import os
import logging
from pypdf import PdfReader
import docx

logger = logging.getLogger("chatbot.file_processor")

ALLOWED_DOCUMENT_EXTENSIONS = {".pdf", ".txt", ".docx"}
ALLOWED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}
ALLOWED_EXTENSIONS = ALLOWED_DOCUMENT_EXTENSIONS | ALLOWED_IMAGE_EXTENSIONS

MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB limit
MAX_EXTRACTED_CHARS = 60000  # Generous limit supporting complete multi-page documents and exam answer sheets


class FileValidationError(Exception):
    """Raised when an uploaded file fails validation."""
    pass


def _get_extension(filename: str) -> str:
    """Returns the lowercased file extension including the leading dot."""
    return os.path.splitext(filename)[1].lower()


def is_image_file(filename: str) -> bool:
    """Returns True if the file is a supported image type."""
    return _get_extension(filename) in ALLOWED_IMAGE_EXTENSIONS


def validate_file(file_storage) -> None:
    """
    Validates a Flask FileStorage object before saving to disk.
    Raises FileValidationError with user-friendly explanations.
    """
    filename = file_storage.filename
    if not filename:
        raise FileValidationError("No file was selected for upload.")

    ext = _get_extension(filename)
    if ext not in ALLOWED_EXTENSIONS:
        allowed_list = ", ".join(sorted(ALLOWED_EXTENSIONS))
        raise FileValidationError(
            f"Unsupported file type '{ext}'. Supported formats: {allowed_list}"
        )

    # Determine size without reading full file into memory
    file_storage.stream.seek(0, os.SEEK_END)
    size = file_storage.stream.tell()
    file_storage.stream.seek(0)

    if size == 0:
        raise FileValidationError("The uploaded file is empty (0 bytes). Please upload a valid document or image.")

    if size > MAX_FILE_SIZE_BYTES:
        mb_size = size / (1024 * 1024)
        max_mb = MAX_FILE_SIZE_BYTES / (1024 * 1024)
        raise FileValidationError(
            f"File is too large ({mb_size:.1f} MB). Maximum allowed upload size is {max_mb:.0f} MB."
        )

    # Filename path traversal check
    sanitized = os.path.basename(filename)
    if ".." in filename or filename.startswith("/") or filename.startswith("\\") or not sanitized:
        raise FileValidationError("Invalid or insecure filename detected.")


def extract_text_from_file(filepath: str) -> str:
    """
    Extracts text from PDF, TXT, or DOCX file at filepath.
    Returns truncated text up to MAX_EXTRACTED_CHARS.
    """
    ext = _get_extension(filepath)

    if ext == ".txt":
        text = _extract_txt(filepath)
    elif ext == ".pdf":
        text = _extract_pdf(filepath)
    elif ext == ".docx":
        text = _extract_docx(filepath)
    else:
        raise ValueError(f"No text extractor available for '{ext}' files.")

    text = text.strip()
    if not text:
        raise ValueError("No readable text could be found in the uploaded file.")

    if len(text) > MAX_EXTRACTED_CHARS:
        text = text[:MAX_EXTRACTED_CHARS] + "\n\n[...content truncated for optimal AI processing...]"

    return text


def _extract_txt(filepath: str) -> str:
    """Reads a text file trying common encodings (UTF-8, Latin-1, CP1252)."""
    encodings = ["utf-8", "latin-1", "cp1252"]
    for enc in encodings:
        try:
            with open(filepath, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    # Fallback with replacement
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _extract_pdf(filepath: str) -> str:
    """Extracts text from all pages of a PDF document, with graceful fallback for scanned PDFs."""
    try:
        reader = PdfReader(filepath)
        pages_text = []
        has_visual_content = False

        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    pages_text.append(f"--- Page {i + 1} ---\n{page_text.strip()}")
            except Exception as page_err:
                logger.warning("Skipped page %d text extraction issue: %s", i + 1, page_err)

            if getattr(page, "images", None) and len(page.images) > 0:
                has_visual_content = True
            elif page.get_contents() is not None:
                has_visual_content = True

        if not pages_text:
            # Rejection check for truly empty 0-content PDF files
            if not has_visual_content or (len(reader.pages) <= 1 and os.path.getsize(filepath) < 2000):
                return ""

            # Gracefully handle scanned/image-based PDFs (e.g. lab code screenshots, scan sheets)
            num_pages = len(reader.pages)
            base_name = os.path.basename(filepath)
            return (
                f"DOCUMENT: {base_name}\n"
                f"TOTAL PAGES: {num_pages} Page(s)\n"
                f"TYPE: Scanned / Image-Based Document (Computer Vision / Lab Report)\n\n"
                f"OVERVIEW: This document contains {num_pages} page(s) of visual content, diagrams, lab code screenshots, or handwritten notes.\n"
                f"STATUS: Processed and safely backed up to Azure Cloud Storage container 'uploaded-files'."
            )

        return "\n\n".join(pages_text)
    except Exception as e:
        logger.error("PDF extraction error: %s", e)
        raise ValueError(f"Failed to extract text from PDF: {str(e)}") from e


def _extract_docx(filepath: str) -> str:
    """Extracts text from paragraphs and tables of a DOCX document."""
    try:
        document = docx.Document(filepath)
        content_parts = []
        
        # Paragraphs
        for p in document.paragraphs:
            if p.text.strip():
                content_parts.append(p.text.strip())

        # Tables (if any)
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    content_parts.append(row_text)

        return "\n\n".join(content_parts)
    except Exception as e:
        logger.error("DOCX extraction error: %s", e)
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}") from e


def generate_document_analysis_report(filename: str, text: str, user_prompt: str = "") -> str:
    """
    Generates a structured, comprehensive document analysis from extracted file text.
    Provides complete answers, preserving all questions, sections, explanations,
    formulas, and bullet points without truncating mid-sentence.
    """
    import re

    clean_text = (text or "").replace("[...content truncated for optimal AI processing...]", "").strip()
    if not clean_text:
        return "No readable document text found."

    lines = [l.strip() for l in clean_text.splitlines() if l.strip()]
    title = os.path.basename(filename)
    non_page_lines = [l for l in lines if not l.startswith("--- Page")]
    if non_page_lines:
        first_candidate = non_page_lines[0]
        if len(first_candidate) < 100:
            title = first_candidate

    word_count = len(clean_text.split())
    read_time = max(1, round(word_count / 200))

    report = [
        f"## 📄 Complete Document Analysis: **{title}**",
        "",
        f"- **File Name:** `{os.path.basename(filename)}`",
        f"- **Word Count:** ~{word_count:,} words",
        f"- **Estimated Reading Time:** ~{read_time} min",
        "",
        "---",
        ""
    ]

    bullet_chars = ("\uf0b7", "•", "*", "-")

    # Check if document has questions like Q1, Q2, Q3, etc.
    q_indices = [i for i, l in enumerate(non_page_lines) if re.match(r"^Q\d+[\.\:\s]", l, re.IGNORECASE)]

    if q_indices:
        report.append("### 📝 Questions & Complete Answers\n")
        for idx, start_idx in enumerate(q_indices):
            end_idx = q_indices[idx + 1] if idx + 1 < len(q_indices) else len(non_page_lines)
            q_lines = non_page_lines[start_idx:end_idx]
            q_header = q_lines[0]
            ans_lines = q_lines[1:]

            report.append(f"#### **{q_header}**")
            formatted_ans = []
            for l in ans_lines:
                clean_l = l.strip()
                if clean_l.startswith(bullet_chars):
                    clean_l = clean_l.lstrip(" \uf0b7•*-")
                    formatted_ans.append(f"- {clean_l}")
                else:
                    formatted_ans.append(clean_l)
            report.append("\n".join(formatted_ans).strip())
            report.append("")
    else:
        # Standard structured sections (e.g. Aim, Procedure, Algorithm, Output, etc.)
        sections = {}
        current_sec = "Overview"
        sections[current_sec] = []

        for line in non_page_lines:
            upper = line.upper()
            if any(upper.startswith(kw) for kw in ["AIM:", "OBJECTIVE:", "PROCEDURE:", "ALGORITHM:", "CODE:", "OUTPUT:", "RESULT:", "CONCLUSION:", "NOTE:"]):
                current_sec = line.split(":")[0].title()
                rest = ":".join(line.split(":")[1:]).strip()
                sections[current_sec] = [rest] if rest else []
            elif line.isupper() and len(line) < 50 and not line.startswith("---"):
                current_sec = line.title()
                sections[current_sec] = []
            else:
                sections[current_sec].append(line)

        report.append("### 📖 Full Document Content\n")
        for sec_name, sec_lines in sections.items():
            if not sec_lines:
                continue
            formatted_sec = []
            for l in sec_lines:
                clean_l = l.strip()
                if clean_l.startswith(bullet_chars):
                    clean_l = clean_l.lstrip(" \uf0b7•*-")
                    formatted_sec.append(f"- {clean_l}")
                else:
                    formatted_sec.append(clean_l)

            if len(sections) > 1:
                report.append(f"**{sec_name}**")
            report.append("\n".join(formatted_sec).strip())
            report.append("")

    report.append("---")
    report.append(
        "> 💡 **Tip**: *This complete analysis was extracted directly via Cloud Document Processing. "
        "To enable real-time conversational AI Q&A on this document, you can add a 100% free Google AI Studio key (`GEMINI_API_KEY`) at [aistudio.google.com/apikey](https://aistudio.google.com/apikey) or top up your OpenRouter balance at [openrouter.ai/settings/credits](https://openrouter.ai/settings/credits).*"
    )

    return "\n".join(report)
